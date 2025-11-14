from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import PyPDF2
import re
import os

# ==================== CONFIGURATION ====================
GROQ_API_KEY = os.getenv('GROQ_API_KEY')  # Set this in your environment variables
MODEL_NAME = "meta-llama/llama-4-scout-17b-16e-instruct"  # Use a model available in your Groq account

# ==================== FILE READER ====================
def read_script_file(file_path):
    """
    Read script from TXT, PDF, or DOCX file
    Returns the text content as a string
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_extension == '.txt':
            # Read plain text file
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        
        elif file_extension == '.pdf':
            # Read PDF file
            text = ""
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text
        
        elif file_extension == '.docx':
            # Read Word document
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        
        else:
            raise ValueError(f"Unsupported file format: {file_extension}. Please use .txt, .pdf, or .docx files.")
    
    except Exception as e:
        raise Exception(f"Error reading file: {str(e)}")

# ==================== SCRIPT PARSER ====================
class ScriptParser:
    """Analyzes screenplay for quantifiable metrics"""
    
    def __init__(self, script_text):
        self.script = script_text
        self.lines = script_text.split('\n')
        
    def count_pages(self):
        """Estimate page count (industry standard: ~55 lines per page)"""
        return len(self.lines) / 55
    
    def check_format_compliance(self):
        """Check screenplay format adherence, excluding page count"""
        total_checks = 0
        passed_checks = 0
        issues = []
        
        # Check 1: Scene headings (INT./EXT.)
        scene_headings = re.findall(r'^(INT\.|EXT\.)', self.script, re.MULTILINE)
        improper_headings = re.findall(r'^(int\.|ext\.|Interior|Exterior)', self.script, re.MULTILINE | re.IGNORECASE)
        total_checks += len(scene_headings) + len(improper_headings)
        passed_checks += len(scene_headings)
        if improper_headings:
            issues.append(f"Found {len(improper_headings)} improperly formatted scene headings (should be INT. or EXT. in caps)")
        
        # Check 2: Character names in CAPS before dialogue
        character_cues = [line for line in self.lines if re.match(r'^\s{10,30}[A-Z\s]+$', line)]
        total_checks += len(character_cues)
        passed_checks += len(character_cues)
        
        # Check 3: Action lines in present tense (basic check)
        past_tense_words = re.findall(r'\b(walked|talked|ran|said|went|came|saw|looked)\b', self.script, re.IGNORECASE)
        total_checks += 1
        if len(past_tense_words) < 10:  # Allow some flexibility
            passed_checks += 1
        else:
            issues.append(f"Found {len(past_tense_words)} potential past-tense verbs (should use present tense)")
        
        score = (passed_checks / total_checks * 100) if total_checks > 0 else 0
        return score, issues
    
    def check_grammar_spelling(self):
        """Basic grammar check (counts common errors)"""
        total_words = len(self.script.split())
        
        errors = 0
        error_examples = []
        
        # Double spaces
        double_spaces = len(re.findall(r'  +', self.script))
        errors += double_spaces
        if double_spaces > 0:
            error_examples.append(f"{double_spaces} instances of double spacing")
        
        # Missing periods at end of sentences
        missing_periods = len(re.findall(r'[a-z]\n[A-Z]', self.script))
        errors += missing_periods
        if missing_periods > 0:
            error_examples.append(f"{missing_periods} potential missing periods")
        
        # Basic misspellings
        common_errors = ['teh', 'recieve', 'occured', 'seperate', 'definately']
        for error in common_errors:
            count = len(re.findall(r'\b' + error + r'\b', self.script, re.IGNORECASE))
            if count > 0:
                errors += count
                error_examples.append(f"'{error}' appears {count} time(s)")
        
        score = max(0, ((total_words - errors) / total_words * 100)) if total_words > 0 else 0
        return score, error_examples
    
    def analyze_dialogue_ratio(self):
        """Calculate dialogue vs action ratio using character name preceding dialogue logic"""
        dialogue_word_count = 0
        reading_dialogue = False
        
        def is_scene_heading(line):
            return bool(re.match(r'^(INT\.|EXT\.)', line.strip()))
        
        def is_character_name_line(line):
            stripped = line.rstrip()
            return bool(re.match(r'^\s{10,30}[A-Z\s]+$', stripped))
        
        for line in self.lines:
            if is_character_name_line(line):
                reading_dialogue = True
                continue
            if reading_dialogue:
                if line.strip() == "" or is_scene_heading(line) or is_character_name_line(line):
                    reading_dialogue = False
                else:
                    dialogue_word_count += len(line.split())
        
        total_word_count = len(self.script.split())
        
        dialogue_ratio = (dialogue_word_count / total_word_count * 100) if total_word_count > 0 else 0
        
        feedback = []
        if dialogue_ratio < 30:
            feedback.append("Script is action-heavy (less than 30% dialogue)")
        elif dialogue_ratio > 50:
            feedback.append("Script is dialogue-heavy (more than 50% dialogue)")
        else:
            feedback.append("Dialogue-to-action ratio is well-balanced (30-50%)")
        
        ideal_ratio = 40
        deviation = abs(dialogue_ratio - ideal_ratio)
        score = max(0, 100 - (deviation * 2))
        
        return score, dialogue_ratio, feedback
    
    def analyze_scene_structure(self):
        """Analyze scene count and length"""
        scene_headings = re.findall(r'^(INT\.|EXT\.)[^\n]+', self.script, re.MULTILINE)
        scene_count = len(scene_headings)
        
        page_count = self.count_pages()
        avg_scene_length = (page_count / scene_count) if scene_count > 0 else 0
        
        feedback = []
        
        if avg_scene_length < 1.0:
            feedback.append(f"Scenes are very short (avg {avg_scene_length:.1f} pages) - may feel choppy")
        elif avg_scene_length > 4.0:
            feedback.append(f"Scenes are quite long (avg {avg_scene_length:.1f} pages) - may slow pacing")
        else:
            feedback.append(f"Scene length is good (avg {avg_scene_length:.1f} pages)")
        
        ideal_scene_count_min = 40
        ideal_scene_count_max = 60
        
        if ideal_scene_count_min <= scene_count <= ideal_scene_count_max:
            score = 100
        elif scene_count < ideal_scene_count_min:
            score = max(0, 100 - (ideal_scene_count_min - scene_count) * 2)
        else:
            score = max(0, 100 - (scene_count - ideal_scene_count_max) * 2)
        
        return score, scene_count, avg_scene_length, feedback
    
    def analyze_white_space(self):
        """Check white space ratio for readability"""
        total_lines = len(self.lines)
        empty_lines = sum(1 for line in self.lines if line.strip() == '')
        
        white_space_ratio = (empty_lines / total_lines * 100) if total_lines > 0 else 0
        
        feedback = []
        if white_space_ratio < 35:
            feedback.append("Script is dense with text (may be hard to read)")
            score = 60
        elif white_space_ratio > 65:
            feedback.append("Script has too much white space (may waste pages)")
            score = 70
        else:
            feedback.append("White space is well-balanced (easy to read)")
            score = 100
        
        return score, white_space_ratio, feedback
    
    def get_all_metrics(self):
        """Run all quantifiable analyses"""
        format_score, format_issues = self.check_format_compliance()
        grammar_score, grammar_issues = self.check_grammar_spelling()
        dialogue_score, dialogue_ratio, dialogue_feedback = self.analyze_dialogue_ratio()
        scene_score, scene_count, avg_scene_length, scene_feedback = self.analyze_scene_structure()
        white_space_score, white_space_ratio, white_space_feedback = self.analyze_white_space()
        
        return {
            'format': {'score': format_score, 'issues': format_issues},
            'grammar': {'score': grammar_score, 'issues': grammar_issues},
            'dialogue': {'score': dialogue_score, 'ratio': dialogue_ratio, 'feedback': dialogue_feedback},
            'scenes': {'score': scene_score, 'count': scene_count, 'avg_length': avg_scene_length, 'feedback': scene_feedback},
            'white_space': {'score': white_space_score, 'ratio': white_space_ratio, 'feedback': white_space_feedback},
            'page_count': self.count_pages()
        }

# ==================== AI EVALUATOR ====================
def get_ai_qualitative_analysis(script_text, metrics):
    """Use Groq AI for qualitative analysis"""
    
    client = Groq(api_key=GROQ_API_KEY)
    
    prompt = f"""You are a professional screenplay analyst. Below is the script and its quantitative metrics.

SCRIPT METRICS:
- Page Count: {metrics['page_count']:.1f}
- Scene Count: {metrics['scenes']['count']}
- Average Scene Length: {metrics['scenes']['avg_length']:.1f} pages
- Dialogue Ratio: {metrics['dialogue']['ratio']:.1f}%

SCRIPT:
{script_text[:3000]}  
[... script continues ...]

Please evaluate the screenplay on each category below and provide:

- A score between 1 and 10
- Specific strengths with quotes or references to the script
- Specific weaknesses with concrete examples

CATEGORIES:

1. CHARACTER DEVELOPMENT
- Are main characters properly motivated and growing?
- Are characters authentic and relatable?
- Provide examples.

2. DIALOGUE QUALITY
- Does dialogue sound natural and distinct for characters?
- Is subtext used instead of exposition?
- Does dialogue advance the plot?
- Provide examples.

3. STORY CONCEPT & ORIGINALITY
- Is the premise original and compelling?
- Does the story have a clear hook?
- Reference specific elements.

OUTPUT FORMAT:
Score: X/10
Strengths: <bullet points>
Weaknesses: <bullet points>

Be specific and concise. Reference exact script lines or scenes.
"""
    
    response = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[{"role": "user", "content": prompt}],
        temperature=0,
        max_tokens=2000,
        seed=42
    )
    
    return response.choices[0].message.content

# ==================== CALCULATE FINAL SCORE ====================
def calculate_final_score(metrics, ai_scores):
    """Weighted scoring system: 70% Objective + 30% Qualitative"""
    
    objective_score = (
        metrics['format']['score'] * 0.15 +
        metrics['grammar']['score'] * 0.10 +
        metrics['dialogue']['score'] * 0.10 +
        metrics['scenes']['score'] * 0.10 +
        metrics['white_space']['score'] * 0.05
    ) * 0.70
    
    character_score = ai_scores.get('character', 7) * 10
    dialogue_quality_score = ai_scores.get('dialogue', 7) * 10
    originality_score = ai_scores.get('originality', 7) * 10
    
    qualitative_score = (
        character_score * 0.15 +
        dialogue_quality_score * 0.10 +
        originality_score * 0.05
    ) * 0.30
    
    final_score = objective_score + qualitative_score
    
    return final_score

# ==================== GENERATE REPORT ====================
def generate_report(script_title, metrics, ai_analysis, final_score):
    doc = Document()
    
    title = doc.add_heading('NEKA Script Evaluation Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('Script Information', level=1)
    doc.add_paragraph(f'Title: {script_title}')
    doc.add_paragraph(f'Page Count: {metrics["page_count"]:.1f} pages')
    doc.add_paragraph(f'Scene Count: {metrics["scenes"]["count"]} scenes')
    doc.add_paragraph('')
    
    score_para = doc.add_heading(f'OVERALL SCORE: {final_score:.1f}/100', level=1)
    score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if final_score >= 80:
        recommendation = "RECOMMEND"
        color = RGBColor(0, 128, 0)
    elif final_score >= 60:
        recommendation = "CONSIDER"
        color = RGBColor(255, 165, 0)
    else:
        recommendation = "PASS"
        color = RGBColor(255, 0, 0)
    
    rec_para = doc.add_paragraph(f'Recommendation: {recommendation}')
    rec_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rec_para.runs[0].font.bold = True
    rec_para.runs[0].font.color.rgb = color
    doc.add_paragraph('')
    
    doc.add_heading('TIER 1: QUANTIFIABLE METRICS (70% weight)', level=1)
    
    doc.add_heading(f'1. Format Compliance: {metrics["format"]["score"]:.1f}/100', level=2)
    for issue in metrics['format']['issues']:
        doc.add_paragraph(f'• {issue}', style='List Bullet')
    doc.add_paragraph('')
    
    doc.add_heading(f'2. Grammar & Spelling: {metrics["grammar"]["score"]:.1f}/100', level=2)
    if metrics['grammar']['issues']:
        for issue in metrics['grammar']['issues']:
            doc.add_paragraph(f'• {issue}', style='List Bullet')
    else:
        doc.add_paragraph('• No major issues detected')
    doc.add_paragraph('')
    
    doc.add_heading(f'3. Dialogue Balance: {metrics["dialogue"]["score"]:.1f}/100', level=2)
    doc.add_paragraph(f'Dialogue comprises {metrics["dialogue"]["ratio"]:.1f}% of the script (Ideal: 35-45%)')
    for feedback in metrics['dialogue']['feedback']:
        doc.add_paragraph(f'• {feedback}', style='List Bullet')
    doc.add_paragraph('')
    
    doc.add_heading(f'4. Scene Structure: {metrics["scenes"]["score"]:.1f}/100', level=2)
    doc.add_paragraph(f'Average scene length: {metrics["scenes"]["avg_length"]:.1f} pages (Ideal: 1.5-3 pages)')
    for feedback in metrics['scenes']['feedback']:
        doc.add_paragraph(f'• {feedback}', style='List Bullet')
    doc.add_paragraph('')
    
    doc.add_heading(f'5. Readability (White Space): {metrics["white_space"]["score"]:.1f}/100', level=2)
    doc.add_paragraph(f'White space ratio: {metrics["white_space"]["ratio"]:.1f}% (Ideal: 40-60%)')
    for feedback in metrics['white_space']['feedback']:
        doc.add_paragraph(f'• {feedback}', style='List Bullet')
    doc.add_paragraph('')
    
    doc.add_heading('TIER 2: QUALITATIVE ANALYSIS (30% weight)', level=1)
    doc.add_paragraph(ai_analysis)
    
    filename = f'{script_title.replace(" ", "_")}_Evaluation.docx'
    
    # Try to save to Downloads folder, fallback to current directory
    try:
        downloads_path = os.path.join(os.path.expanduser('~'), 'Downloads')
        full_path = os.path.join(downloads_path, filename)
        doc.save(full_path)
        print(f"Report saved to Downloads: {filename}")
    except:
        # Fallback to current directory
        doc.save(filename)
        print(f"Report saved as: {filename}")

# ==================== MAIN EXECUTION ====================
def main():
    """Main execution function"""
    
    print("=== NEKA Script Evaluation System ===\n")
    print("Supported file formats: .txt, .pdf, .docx")
    
    script_file = input("Enter script file path (e.g., script.txt, script.pdf, script.docx): ").strip()
    
    if not os.path.exists(script_file):
        print(f"❌ Error: {script_file} not found!")
        return
    
    try:
        print("\n📄 Reading script file...")
        script_text = read_script_file(script_file)
        
        if not script_text or len(script_text.strip()) == 0:
            print("❌ Error: File is empty or could not be read!")
            return
        
        print(f"✓ Successfully read {len(script_text)} characters")
        
    except Exception as e:
        print(f"❌ Error reading file: {str(e)}")
        return
    
    script_title = input("\nEnter script title: ") or "Untitled Script"
    
    print("\n🔍 Analyzing script quantitative metrics...")
    
    parser = ScriptParser(script_text)
    metrics = parser.get_all_metrics()
    
    print(f"✓ Page count: {metrics['page_count']:.1f}")
    print(f"✓ Scene count: {metrics['scenes']['count']}")
    print(f"✓ Format score: {metrics['format']['score']:.1f}/100")
    print(f"✓ Grammar score: {metrics['grammar']['score']:.1f}/100")
    
    print("\n🤖 Getting AI qualitative analysis...")
    ai_analysis = get_ai_qualitative_analysis(script_text, metrics)
    
    ai_scores = {'character': 7, 'dialogue': 7, 'originality': 7}
    
    print("\n📊 Calculating final score...")
    final_score = calculate_final_score(metrics, ai_scores)
    
    print(f"\n⭐ FINAL SCORE: {final_score:.1f}/100")
    
    print("\n📝 Generating report...")
    generate_report(script_title, metrics, ai_analysis, final_score)
    
    print("\n✅ Evaluation complete!")

if __name__ == "__main__":
    main()
